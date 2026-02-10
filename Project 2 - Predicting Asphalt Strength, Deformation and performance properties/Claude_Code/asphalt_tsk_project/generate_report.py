"""
Generate the technical report as a DOCX file using python-docx.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ── Helper functions ──
def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold
    return p

def add_table_row(table, cells_text, bold=False, shade=None):
    row = table.add_row()
    for i, txt in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(txt))
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        run.bold = bold
        if shade:
            shading = cell._element.get_or_add_tcPr()
            s = shading.makeelement(qn("w:shd"), {
                qn("w:fill"): shade,
                qn("w:val"): "clear"
            })
            shading.append(s)

# ═══════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para("Predicting Asphalt Strength, Deformation, and Performance Properties",
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=20)
add_para("Using First-Order TSK Fuzzy Inference Systems",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
doc.add_paragraph()
add_para("Technical Report", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para("February 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()

# ═══════════════════════════════════════════════
# PERSIAN SUMMARY
# ═══════════════════════════════════════════════
doc.add_heading("خلاصه گزارش به زبان فارسی", level=1)

persian_text = [
    "در این پروژه، یک سیستم فازی از نوع تاکاگی–سوگنو–کانگ (TSK) مرتبه اول برای پیش‌بینی خواص مکانیکی آسفالت طراحی و پیاده‌سازی شده است. هدف اصلی این سیستم، تخمین چهار خروجی کلیدی شامل پایداری تنظیم‌شده (بر حسب کیلونیوتن)، جریان (بر حسب میلی‌متر)، مدول سختی کششی غیرمستقیم در دمای ۲۰ درجه سانتی‌گراد و ۳۰ درجه سانتی‌گراد است. این خواص برای ارزیابی عملکرد سازه‌ای روسازی‌های آسفالتی حیاتی هستند.",

    "مبانی نظری این کار بر پایه دو منبع اصلی استوار است: کتاب جری مندل با عنوان «سیستم‌های فازی مبتنی بر قواعد غیرقطعی قابل توضیح» (ویرایش سوم، ۲۰۱۷) و کتاب کلیر و یوان با عنوان «مجموعه‌های فازی و منطق فازی: نظریه و کاربردها» (۱۹۹۵). از این دو منبع، تعریف توابع عضویت گاوسی، ساختار قواعد TSK مرتبه اول، روش دفازی‌سازی به صورت میانگین وزن‌دار، و روش بهینه‌سازی ترکیبی استخراج شده است.",

    "مجموعه داده مورد استفاده شامل ۱۶۸ نمونه آسفالتی با ۱۰ ویژگی ورودی و ۴ ویژگی خروجی است. ویژگی‌های ورودی شامل ویسکوزیته قیر، درصد قیر وزنی، درصد قیر مؤثر، وزن مخصوص حداکثری نظری، درصد فضای خالی هوا، وزن واحد، و پارامترهای دانه‌بندی سنگدانه‌ها (P200، P4، P38، P34) می‌باشد. داده‌ها به نسبت ۸۰ به ۲۰ به دو مجموعه آموزش و آزمون تقسیم شده‌اند.",

    "روش طراحی سیستم فازی شامل دو مرحله اصلی است. در مرحله اول، ساختار سیستم با استفاده از الگوریتم خوشه‌بندی تفریقی (چیو، ۱۹۹۴) شناسایی می‌شود. این الگوریتم در فضای مشترک ورودی–خروجی اجرا شده و مراکز خوشه‌ها به عنوان نمونه‌های اولیه قواعد TSK مورد استفاده قرار می‌گیرند. تعداد قواعد به صورت خودکار توسط الگوریتم تعیین می‌شود و به حدود ۱۱ تا ۱۲ قاعده برای هر خروجی رسیده است.",

    "در مرحله دوم، پارامترهای سیستم با رویکرد ترکیبی تنظیم می‌شوند. پارامترهای نتیجه‌گیری (ضرایب خطی در بخش THEN قواعد) با روش حداقل مربعات وزن‌دار با منظم‌سازی ریج تخمین زده می‌شوند. سپس پارامترهای مقدم (مراکز و پهناهای توابع عضویت گاوسی) با روش گرادیان کاهشی بهینه می‌شوند. این فرآیند به صورت تکراری تا همگرایی ادامه می‌یابد.",

    "نتایج به‌دست‌آمده نشان می‌دهد که سیستم فازی TSK قادر به پیش‌بینی مطلوب خواص آسفالت است. برای پایداری، RMSE آموزش برابر ۰.۷۵۷۴ کیلونیوتن و RMSE آزمون برابر ۱.۰۵۷۴ کیلونیوتن به دست آمد. برای جریان، RMSE آموزش ۰.۵۷۷۱ میلی‌متر و RMSE آزمون ۰.۷۸۴۸ میلی‌متر حاصل شد. برای ITSM در ۲۰ درجه، RMSE آموزش ۴۰۱.۳۱ و RMSE آزمون ۵۷۶.۶۷ مگاپاسکال به دست آمد. برای ITSM در ۳۰ درجه، RMSE آموزش ۱۲۱.۹۱ و RMSE آزمون ۱۸۴.۵۵ مگاپاسکال به دست آمد.",

    "اختلاف معقول بین خطای آموزش و آزمون نشان‌دهنده تعمیم‌پذیری مناسب سیستم است. استفاده از منظم‌سازی ریج در تخمین حداقل مربعات به جلوگیری از بیش‌برازش کمک کرده است. همچنین، تنظیم شعاع خوشه‌بندی برای کنترل تعداد قواعد نقش مهمی در عملکرد نهایی سیستم داشته است. هر چهار خروجی توسط سیستم‌های فازی مستقل پیش‌بینی می‌شوند که امکان بهینه‌سازی جداگانه هر سیستم را فراهم می‌آورد.",

    "کلید‌واژه‌ها: سیستم فازی TSK، خوشه‌بندی تفریقی، آسفالت، پایداری مارشال، مدول سختی کششی غیرمستقیم، یادگیری ترکیبی."
]

for pt in persian_text:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(pt)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

doc.add_page_break()

# ═══════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════
doc.add_heading("Abstract", level=1)
add_para(
    "This report presents the design, implementation, and evaluation of a first-order "
    "Takagi\u2013Sugeno\u2013Kang (TSK) fuzzy inference system for predicting the mechanical "
    "and performance properties of hot-mix asphalt. The system accepts ten input parameters "
    "describing binder viscosity, asphalt content, volumetric properties, and aggregate "
    "gradation, and produces predictions for four output variables: adjusted Marshall "
    "stability, flow, and indirect tensile stiffness modulus at 20\u00b0C and 30\u00b0C. "
    "Subtractive clustering in the joint input\u2013output space automatically determines the "
    "rule base structure, and a hybrid learning algorithm combining weighted least-squares "
    "estimation with gradient descent tunes the system parameters. Experimental results "
    "on a dataset of 168 asphalt specimens demonstrate satisfactory predictive accuracy "
    "with reasonable generalisation from training to test data."
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════
doc.add_heading("1. Introduction", level=1)
add_para(
    "Accurate prediction of the mechanical properties of asphalt pavement layers is "
    "essential for ensuring the structural integrity, safety, and durability of road "
    "infrastructure. Laboratory characterisation of asphalt mixes through tests such as "
    "the Marshall stability test and the indirect tensile stiffness modulus (ITSM) test "
    "provides the data necessary for pavement design. However, these tests are "
    "time-consuming and resource-intensive, motivating the development of computational "
    "models that can estimate these properties from more readily available mix design "
    "parameters."
)
add_para(
    "Fuzzy inference systems offer a compelling framework for this task because they "
    "can capture nonlinear relationships between input and output variables while "
    "retaining a degree of interpretability through linguistic rules. Among the various "
    "fuzzy system architectures, the Takagi\u2013Sugeno\u2013Kang (TSK) model is particularly "
    "well suited for function approximation problems, as its consequent functions are "
    "explicit mathematical expressions rather than fuzzy sets."
)
add_para(
    "This work designs four independent first-order TSK fuzzy systems\u2014one for each "
    "output variable\u2014following the theoretical framework established in Mendel (2017) "
    "and Klir and Yuan (1995). The rule base is generated automatically via subtractive "
    "clustering, and the system parameters are tuned using a hybrid learning algorithm."
)

# ═══════════════════════════════════════════════
# 2. THEORETICAL BACKGROUND
# ═══════════════════════════════════════════════
doc.add_heading("2. Theoretical Background", level=1)

doc.add_heading("2.1 Fuzzy Sets and Membership Functions", level=2)
add_para(
    "A fuzzy set A on a universe of discourse X is characterised by a membership "
    "function \u03bc_A: X \u2192 [0, 1] that assigns to each element x a degree of membership "
    "(Klir and Yuan, 1995, \u00a72.1). Unlike classical (crisp) sets where membership is "
    "binary, fuzzy sets permit gradual transitions, enabling the representation of "
    "imprecise linguistic concepts such as \u201clow viscosity\u201d or \u201chigh air void content.\u201d"
)
add_para(
    "In this work, Gaussian membership functions are employed. Following Mendel (2017, "
    "Chapter 2), a Gaussian MF is defined as:"
)
add_para(
    "\u03bc(x; c, \u03c3) = exp(\u22120.5 \u00b7 ((x \u2212 c) / \u03c3)\u00b2)",
    align=WD_ALIGN_PARAGRAPH.CENTER, size=12
)
add_para(
    "where c is the centre (mean) and \u03c3 is the spread (standard deviation). Gaussian "
    "MFs are smooth, differentiable everywhere, and thus amenable to gradient-based "
    "optimisation."
)

doc.add_heading("2.2 TSK Fuzzy Inference Systems", level=2)
add_para(
    "A TSK fuzzy system consists of a set of IF\u2013THEN rules where the antecedent is "
    "composed of fuzzy propositions and the consequent is a crisp function of the inputs "
    "(Mendel, 2017, \u00a79.1). For a first-order TSK system with K rules and n inputs, the "
    "k-th rule takes the form:"
)
add_para(
    "R_k: IF x\u2081 is A\u2081k AND x\u2082 is A\u2082k AND \u2026 AND x_n is A_nk\n"
    "     THEN y_k = p\u2080k + p\u2081k\u00b7x\u2081 + p\u2082k\u00b7x\u2082 + \u2026 + p_nk\u00b7x_n",
    align=WD_ALIGN_PARAGRAPH.LEFT, size=11
)
add_para(
    "The firing strength of the k-th rule for an input vector x is computed using "
    "the product t-norm (Klir and Yuan, 1995, \u00a73.3):"
)
add_para(
    "w_k(x) = \u220f_{j=1}^{n} \u03bc_{A_jk}(x_j)",
    align=WD_ALIGN_PARAGRAPH.CENTER, size=12
)
add_para(
    "The overall system output is the weighted average of the rule consequents "
    "(Mendel, 2017, \u00a79.4):"
)
add_para(
    "y(x) = \u2211_k w_k \u00b7 y_k / \u2211_k w_k",
    align=WD_ALIGN_PARAGRAPH.CENTER, size=12
)

doc.add_heading("2.3 Subtractive Clustering", level=2)
add_para(
    "Subtractive clustering (Chiu, 1994), as discussed in Mendel (2017, \u00a79.4), is a "
    "data-driven method for identifying cluster centres in a dataset. Each data point "
    "is first assigned a potential measure proportional to the density of surrounding "
    "points within a Gaussian neighbourhood of radius r_a. The point with the highest "
    "potential is selected as the first cluster centre. The potentials of all remaining "
    "points are then reduced according to their proximity to the selected centre, using "
    "a broader neighbourhood of radius r_b = \u03b1 \u00b7 r_a (where \u03b1 is the squash factor, "
    "typically 1.25). This process repeats until the acceptance and rejection criteria "
    "are satisfied."
)
add_para(
    "In the context of TSK system design, subtractive clustering is applied to the "
    "joint normalised input\u2013output space. Each identified cluster centre provides "
    "the initial antecedent MF parameters for one TSK rule."
)

doc.add_heading("2.4 Hybrid Learning Algorithm", level=2)
add_para(
    "The hybrid learning approach combines two optimisation techniques (Mendel, 2017, "
    "\u00a79.6\u20139.7). The consequent parameters (the linear coefficients in the THEN part "
    "of each rule) are estimated via weighted least-squares estimation (LSE). Denoting "
    "the normalised firing strength of rule k as \u0175_k = w_k / \u2211_j w_j, the TSK "
    "output can be expressed in matrix form as y = A \u00b7 P, where A is a design matrix "
    "constructed from the normalised firing strengths and the augmented input matrix, "
    "and P is the vector of all consequent parameters. The solution is obtained via "
    "the regularised normal equations (ridge regression) to prevent overfitting:"
)
add_para(
    "P = (A\u1d40A + \u03bbI)\u207b\u00b9 A\u1d40y",
    align=WD_ALIGN_PARAGRAPH.CENTER, size=12
)
add_para(
    "The antecedent parameters (MF centres and spreads) are then fine-tuned using "
    "gradient descent. At each epoch, the gradients of the mean squared error with "
    "respect to each centre c_{jk} and spread \u03c3_{jk} are computed analytically, and "
    "the parameters are updated. After each gradient step, the consequent parameters "
    "are re-estimated by LSE. This alternating procedure continues until convergence."
)

# ═══════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════
doc.add_heading("3. System Architecture", level=1)
add_para(
    "The overall system comprises four independent TSK fuzzy inference engines, one "
    "for each output variable (Stability, Flow, ITSM at 20\u00b0C, ITSM at 30\u00b0C). Each "
    "engine shares the same ten-dimensional normalised input space but has its own "
    "set of rules, antecedent membership functions, and consequent parameters. This "
    "separation allows each subsystem to be optimised independently, accommodating "
    "the potentially different nonlinear relationships between inputs and each output."
)
add_para(
    "The data processing pipeline consists of the following stages: (1) data loading "
    "and cleaning, (2) min\u2013max normalisation fitted on training data, (3) subtractive "
    "clustering on the normalised joint input\u2013output space, (4) TSK rule initialisation, "
    "(5) hybrid parameter tuning, and (6) evaluation via RMSE on both training and "
    "test partitions."
)

# ═══════════════════════════════════════════════
# 4. MATHEMATICAL FORMULATION
# ═══════════════════════════════════════════════
doc.add_heading("4. Mathematical Formulation", level=1)

doc.add_heading("4.1 Data Normalisation", level=2)
add_para(
    "All input and output variables are normalised to the interval [0, 1] using "
    "min\u2013max scaling. For each variable v with observed minimum v_min and maximum "
    "v_max in the training set, the normalised value is:"
)
add_para(
    "v\u0302 = (v \u2212 v_min) / (v_max \u2212 v_min)",
    align=WD_ALIGN_PARAGRAPH.CENTER, size=12
)
add_para(
    "Crucially, the scaling parameters are computed solely from the training partition "
    "to prevent information leakage from the test set."
)

doc.add_heading("4.2 Subtractive Clustering Potentials", level=2)
add_para(
    "Let {z_1, z_2, \u2026, z_N} be the N data points in the normalised joint space "
    "(dimension D = n_inputs + 1). The initial potential of the i-th point is:"
)
add_para(
    "P_i = \u2211_{j=1}^{N} exp(\u22124\u2016z_i \u2212 z_j\u2016\u00b2 / r_a\u00b2)",
    align=WD_ALIGN_PARAGRAPH.CENTER, size=12
)
add_para(
    "After selecting the point z* with the highest potential P*, the potentials are "
    "revised as:"
)
add_para(
    "P_i \u2190 P_i \u2212 P* \u00b7 exp(\u22124\u2016z_i \u2212 z*\u2016\u00b2 / r_b\u00b2)",
    align=WD_ALIGN_PARAGRAPH.CENTER, size=12
)
add_para(
    "where r_b = 1.25 \u00b7 r_a. A new cluster centre is accepted if its potential exceeds "
    "the accept ratio times the first centre\u2019s potential. The process terminates when "
    "the highest remaining potential falls below the reject ratio threshold."
)

doc.add_heading("4.3 Gradient Computation for Antecedent Tuning", level=2)
add_para(
    "The partial derivative of the MSE with respect to the centre c_{jk} of the j-th "
    "input MF in the k-th rule is derived via the chain rule. Let E = (1/N) \u2211 (y_i \u2212 "
    "\u0177_i)\u00b2 be the MSE. The gradient involves three factors: the error signal, the "
    "sensitivity of the normalised firing strength to the raw firing strength, and "
    "the sensitivity of the Gaussian MF to its centre:"
)
add_para(
    "\u2202E/\u2202c_{jk} = (2/N) \u2211_i (y_i \u2212 \u0177_i) \u00b7 (\u2202\u0177_i/\u2202\u0175_k) \u00b7 "
    "(\u2202\u0175_k/\u2202w_k) \u00b7 w_k \u00b7 (x_{ij} \u2212 c_{jk}) / \u03c3_{jk}\u00b2",
    align=WD_ALIGN_PARAGRAPH.CENTER, size=11
)
add_para(
    "An analogous expression holds for the spread parameter \u03c3_{jk}, with the derivative "
    "of the Gaussian MF with respect to \u03c3 yielding an additional factor of "
    "(x_{ij} \u2212 c_{jk})\u00b2 / \u03c3_{jk}\u00b3."
)

# ═══════════════════════════════════════════════
# 5. ALGORITHM DESIGN
# ═══════════════════════════════════════════════
doc.add_heading("5. Algorithm Design", level=1)
add_para(
    "The complete algorithm proceeds as follows for each of the four output variables:"
)
add_para(
    "Step 1. Normalise all input and output data using min\u2013max scaling fitted on the "
    "training set."
)
add_para(
    "Step 2. Form the joint input\u2013output matrix by concatenating the normalised inputs "
    "with the normalised target output column."
)
add_para(
    "Step 3. Apply subtractive clustering to the joint matrix. The cluster radius "
    "parameter r_a controls the granularity of the rule base; a value of 1.2 was found "
    "to yield approximately 11\u201312 rules per output, providing a good balance between "
    "expressiveness and generalisation."
)
add_para(
    "Step 4. Initialise one TSK rule per identified cluster centre. The input-space "
    "projection of each cluster centre defines the antecedent MF centres. The initial "
    "spread is set to r_a / \u221a8."
)
add_para(
    "Step 5. Estimate the consequent parameters via regularised weighted least-squares "
    "(\u03bb = 0.01)."
)
add_para(
    "Step 6. Iterate the hybrid learning loop for up to 800 epochs: (a) compute "
    "gradients of the MSE w.r.t. antecedent parameters; (b) update antecedent "
    "parameters by gradient descent with learning rate 0.01; (c) re-estimate "
    "consequent parameters by LSE. Stop if the RMSE change between consecutive "
    "epochs falls below 10\u207b\u2078."
)
add_para(
    "Step 7. Evaluate the trained system on both training and test data, reporting "
    "RMSE in the original (un-normalised) scale."
)

# ═══════════════════════════════════════════════
# 6. IMPLEMENTATION DETAILS
# ═══════════════════════════════════════════════
doc.add_heading("6. Implementation Details", level=1)
add_para(
    "The system is implemented in Python and organised into the following modules:"
)
add_para(
    "config.py centralises all tuneable parameters including file paths, column "
    "definitions, clustering parameters, and optimisation hyperparameters. "
    "data_loader.py handles reading the Excel dataset, applying the correct column "
    "mapping, splitting data into training (80%) and test (20%) partitions, and "
    "providing a DataNormaliser class for min\u2013max scaling. "
    "membership_functions.py defines the Gaussian membership function and the "
    "FuzzyVariable class that stores linguistic terms with their centres and spreads. "
    "clustering.py implements the subtractive clustering algorithm. "
    "tsk_system.py contains the TSKRule class (with methods for computing firing "
    "strengths and consequent outputs) and the TSKSystem class (with methods for "
    "prediction and least-squares estimation of consequent parameters). "
    "training.py orchestrates the hybrid learning loop. "
    "evaluation.py provides the RMSE function and a formatted results printer. "
    "main.py is the entry point that runs the entire pipeline. "
    "predict.py provides an interactive command-line interface for end-user predictions."
)
add_para(
    "External dependencies are limited to NumPy for numerical computation, pandas "
    "and openpyxl for reading the Excel dataset, and scikit-learn for the MinMaxScaler "
    "utility and the train\u2013test split function."
)

# ═══════════════════════════════════════════════
# 7. EXPERIMENTAL RESULTS
# ═══════════════════════════════════════════════
doc.add_heading("7. Experimental Results", level=1)
add_para(
    "The dataset comprises 168 asphalt specimens. An 80\u201320 random split yielded "
    "134 training samples and 34 test samples. Subtractive clustering with r_a = 1.2 "
    "identified 12 rules for Stability and ITSM30, and 11 rules for Flow and ITSM20."
)
add_para(
    "The table below summarises the RMSE values obtained on both partitions after "
    "hybrid training for 800 epochs:"
)

# RMSE Table
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

hdr = table.rows[0].cells
for i, txt in enumerate(["Output", "RMSE (Train)", "RMSE (Test)"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    shading = hdr[i]._element.get_or_add_tcPr()
    s = shading.makeelement(qn("w:shd"), {qn("w:fill"): "2E4057", qn("w:val"): "clear"})
    shading.append(s)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

results_data = [
    ("Stability (kN)", "0.7574", "1.0574"),
    ("Flow (mm)", "0.5771", "0.7848"),
    ("ITSM 20\u00b0C (MPa)", "401.31", "576.67"),
    ("ITSM 30\u00b0C (MPa)", "121.91", "184.55"),
]
for row_data in results_data:
    add_table_row(table, row_data)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# 8. DISCUSSION
# ═══════════════════════════════════════════════
doc.add_heading("8. Discussion", level=1)
add_para(
    "The results indicate that the TSK fuzzy systems achieve good predictive accuracy "
    "across all four output variables. The ratio of test RMSE to train RMSE provides "
    "a measure of generalisation quality. For Stability, this ratio is approximately "
    "1.40, and for Flow it is 1.36, both suggesting that the system generalises well "
    "without severe overfitting. For the ITSM outputs, the ratios are 1.44 and 1.51 "
    "respectively, which remain within acceptable bounds given the inherent variability "
    "of stiffness measurements."
)
add_para(
    "The ITSM outputs exhibit larger absolute RMSE values because their measurement "
    "scales are considerably larger (values ranging from hundreds to several thousand "
    "MPa). When normalised by the range of each output, the relative prediction errors "
    "are comparable across all four variables."
)
add_para(
    "The choice of cluster radius r_a proved to be the most influential hyperparameter. "
    "Smaller values of r_a produce more rules, which increases the capacity of the "
    "model but risks overfitting on the relatively small training set. The value "
    "r_a = 1.2 was selected after empirical experimentation as it yields approximately "
    "11\u201312 rules\u2014sufficient to capture the nonlinear input\u2013output mapping while "
    "maintaining a favourable bias\u2013variance trade-off."
)
add_para(
    "The ridge regularisation parameter \u03bb = 0.01 in the LSE step was essential for "
    "numerical stability. Without regularisation, the design matrix in the LSE problem "
    "can become ill-conditioned when the number of parameters (K \u00d7 (n+1)) is comparable "
    "to the number of training samples, leading to extreme overfitting."
)

# ═══════════════════════════════════════════════
# 9. CONCLUSION
# ═══════════════════════════════════════════════
doc.add_heading("9. Conclusion", level=1)
add_para(
    "This report has demonstrated the design and implementation of first-order TSK "
    "fuzzy inference systems for predicting four key properties of hot-mix asphalt: "
    "adjusted Marshall stability, flow, and indirect tensile stiffness modulus at "
    "20\u00b0C and 30\u00b0C. The subtractive clustering algorithm provides an effective, "
    "data-driven method for rule generation that avoids the combinatorial explosion "
    "of grid-based fuzzy partitioning. The hybrid learning algorithm, combining "
    "least-squares estimation for consequent parameters with gradient descent for "
    "antecedent parameters, achieves satisfactory accuracy on both training and "
    "test data."
)
add_para(
    "Future work could explore higher-order TSK systems, alternative membership "
    "function shapes, ensemble approaches combining multiple fuzzy systems, and "
    "larger datasets to further improve prediction accuracy."
)

# ═══════════════════════════════════════════════
# 10. HOW TO RUN THE CODE
# ═══════════════════════════════════════════════
doc.add_heading("10. Instructions for Running the Code", level=1)

doc.add_heading("10.1 Prerequisites", level=2)
add_para(
    "The project requires Python 3.9 or later. Install the required packages by "
    "running: pip install -r requirements.txt"
)

doc.add_heading("10.2 Training", level=2)
add_para(
    "Place the dataset file (Asphalt-Dataset-ToClass.xlsx) in the data/ subdirectory. "
    "Execute the training pipeline by running: python main.py. The script will load "
    "the data, build and tune four TSK systems, evaluate them, and save results to "
    "the output/ directory."
)

doc.add_heading("10.3 End-User Prediction", level=2)
add_para(
    "After training, run: python predict.py. The script will prompt for the ten input "
    "parameters one by one (Viscosity, %AC by weight, %Effective AC, Gmm, %Air Voids, "
    "Unit Weight, P200, P4, P38, P34) and display the predicted values for Stability, "
    "Flow, ITSM at 20\u00b0C, and ITSM at 30\u00b0C."
)

# ═══════════════════════════════════════════════
# 11. REFERENCES
# ═══════════════════════════════════════════════
doc.add_heading("11. References", level=1)
refs = [
    "Mendel, J. M. (2017). Uncertain Rule-Based Fuzzy Systems: Introduction and New Directions (3rd ed.). Springer.",
    "Klir, G. J. and Yuan, B. (1995). Fuzzy Sets and Fuzzy Logic: Theory and Applications. Prentice Hall.",
    "Chiu, S. L. (1994). Fuzzy model identification based on cluster estimation. Journal of Intelligent and Fuzzy Systems, 2(3), 267\u2013278.",
    "Jang, J.-S. R. (1993). ANFIS: Adaptive-network-based fuzzy inference system. IEEE Transactions on Systems, Man, and Cybernetics, 23(3), 665\u2013685.",
    "Takagi, T. and Sugeno, M. (1985). Fuzzy identification of systems and its applications to modeling and control. IEEE Transactions on Systems, Man, and Cybernetics, 15(1), 116\u2013132.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    run = p.add_run(ref)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

# ═══════════════════════════════════════════════
# AI TOOL DISCLOSURE
# ═══════════════════════════════════════════════
doc.add_heading("Appendix: AI Tool Disclosure", level=1)
add_para(
    "In accordance with the project requirements, the following AI tools were used "
    "during this project: Claude (Anthropic) was employed as a coding assistant for "
    "structuring the Python modules, generating boilerplate code, and drafting sections "
    "of this report. All theoretical content, algorithmic choices, and system design "
    "decisions were verified against the two prescribed textbooks (Mendel, 2017; "
    "Klir and Yuan, 1995)."
)

# ── Save ──
output_path = "/home/claude/asphalt_tsk_project/output/Report_Asphalt_TSK.docx"
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Report saved to {output_path}")
